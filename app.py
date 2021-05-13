# СОЗДАНИЕ ПОРТФОЛИО О СЕБЕ (резюме)

from docx import Document
from docx.shared import Inches # Позволить нам установить картинку по желанию размер
import pyttsx3 # Текстовая речь

# text to speech
def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture/Вставка файла в документ и установка размера изоображения
document.add_picture(
    'avatar.png', 
    width=Inches(2.0)
)

# name = 'Max'
# phone_number = '0000000'
# email = 'hello@anymail.com'

# name phone number and email details/
name = input('What is your name? ')
speak('Hello' + name + ' how are you today?')

speak('What is your phone number?')
phone_number = input('What is your phone number? ')

speak('What is your email?')
email = input('What is your email? ')

# about me/добавление описании о себе
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself? '))

# paragraph/Создание параграфа
document.add_paragraph(name + ' | ' + phone_number + ' | ' + email) # написание или вставка текста в документ

# work experience/опыт работы
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Decribe your experience at ' + company + ' ')
p.add_run(experience_details)

# more experience/дополнительный опыт работы
while True:
    has_more_experience = input('Do you have more experience? Yes or No')
    if has_more_experience.lower() == 'yes': # lower облегчить ввод в консоле (могут написать с большой или с маленькой и т.д)
        # copy from work experience
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Decribe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# skills/навыки
document.add_heading('Skills')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do u have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet' 
    else:
        break

# footer/нижний колонтитул
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Intuit QuickBooks course project'

document.save('cv.docx')